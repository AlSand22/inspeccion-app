import streamlit as st
import anthropic
import base64
import json
from io import BytesIO
from datetime import date
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

st.set_page_config(page_title="Inspección Técnica", page_icon="🏗️", layout="centered")

st.markdown("""
<style>
.titulo-negro  {background:#000;color:#fff;padding:10px 16px;border-radius:6px;font-weight:bold;font-size:17px;margin-bottom:8px;}
.titulo-morado {background:#8563E1;color:#fff;padding:8px 16px;border-radius:6px;font-weight:bold;font-size:14px;margin:6px 0;}
.badge-Rechazado   {background:#1a1a1a;color:#fff;padding:3px 10px;border-radius:4px;font-size:12px;font-weight:bold;display:inline-block;}
.badge-Observado   {background:#E07B39;color:#fff;padding:3px 10px;border-radius:4px;font-size:12px;font-weight:bold;display:inline-block;}
.badge-Informativo {background:#7B5EA7;color:#fff;padding:3px 10px;border-radius:4px;font-size:12px;font-weight:bold;display:inline-block;}
.badge-Validado    {background:#9B7FBF;color:#fff;padding:3px 10px;border-radius:4px;font-size:12px;font-weight:bold;display:inline-block;}
.badge-NA          {background:#888;color:#fff;padding:3px 10px;border-radius:4px;font-size:12px;font-weight:bold;display:inline-block;}
</style>
""", unsafe_allow_html=True)

st.markdown('<div class="titulo-negro">🏗️ APP DE INSPECCIÓN TÉCNICA DE INMUEBLES</div>', unsafe_allow_html=True)

# ── API Key en sidebar ────────────────────────────────────────────────────────
with st.sidebar:
    st.markdown("### ⚙️ Configuración")
    # Intentar leer desde Streamlit Secrets primero
    api_key_secret = st.secrets.get("ANTHROPIC_API_KEY", "") if hasattr(st, "secrets") else ""
    if api_key_secret:
        api_key = api_key_secret
        st.success("✅ API Key configurada")
    else:
        api_key = st.text_input("API Key de Anthropic", type="password",
                                 help="O configúrala en Streamlit Secrets")
    st.markdown("---")
    st.markdown("### Tolerancias")
    st.markdown("""
⬛ **Rechazado** — Sustitución completa  
🟧 **Observado** — Reparable  
🟪 **Informativo** — Dentro de tolerancia  
🟣 **Validado** — Sin observaciones  
⬜ **N/A** — No aplicable
    """)

TOLERANCIAS = ["Rechazado", "Observado", "Informativo", "Validado", "N/A"]
AMBIENTES   = ["Dormitorio principal","Dormitorio secundario","Baño principal",
               "Baño secundario","Sala","Comedor","Cocina","Lavandería",
               "Estudio","Pasadizo","Puerta de ingreso","Otro"]

# ── Estado global ─────────────────────────────────────────────────────────────
if "observaciones" not in st.session_state:
    st.session_state.observaciones = []  # lista de dicts
if "upload_key" not in st.session_state:
    st.session_state.upload_key = 0

# ── DATOS DEL INFORME ─────────────────────────────────────────────────────────
st.subheader("Datos del Informe")
c1, c2 = st.columns(2)
with c1:
    propietario  = st.text_input("Propietario")
    telefono     = st.text_input("Teléfono")
    inmobiliaria = st.text_input("Inmobiliaria")
    proyecto     = st.text_input("Proyecto")
with c2:
    direccion    = st.text_input("Dirección")
    nro_depto    = st.text_input("N° Departamento")
    metraje      = st.text_input("Metraje (m²)")
    fecha_insp   = st.date_input("Fecha", value=date.today())

st.divider()

# ── REGISTRAR OBSERVACIÓN ─────────────────────────────────────────────────────
st.markdown('<div class="titulo-negro">📸 REGISTRAR OBSERVACIÓN</div>', unsafe_allow_html=True)

c1, c2 = st.columns(2)
with c1:
    ambiente_sel = st.selectbox("Ambiente", AMBIENTES)
    if ambiente_sel == "Otro":
        ambiente_sel = st.text_input("Especifica el ambiente")
with c2:
    elemento_input = st.text_input("Elemento inspeccionado",
                                    placeholder="Ej: Roseta de ducha, Zócalo, Piso...")

st.markdown("**📷 Fotos de la observación**")
st.caption("Sube primero la foto panorámica y luego las fotos de detalle del defecto.")

fotos = st.file_uploader(
    "Selecciona todas las fotos de esta observación (panorámica + detalle)",
    type=["jpg","jpeg","png"],
    accept_multiple_files=True,
    key=f"fotos_{st.session_state.upload_key}"
)

if fotos:
    cols = st.columns(min(len(fotos), 4))
    for i, f in enumerate(fotos):
        with cols[i % 4]:
            st.image(f, caption=f"{'🌐 Panorámica' if i==0 else f'🔍 Detalle {i}'}", use_container_width=True)

# ── Compresión automática de imágenes ────────────────────────────────────────
def comprimir_imagen(file_bytes, max_bytes=4_000_000):
    from PIL import Image
    import io
    img = Image.open(io.BytesIO(file_bytes))
    if img.mode in ("RGBA", "P"):
        img = img.convert("RGB")
    if max(img.size) > 1920:
        img.thumbnail((1920, 1920), Image.LANCZOS)
    quality = 85
    while quality >= 30:
        buf = io.BytesIO()
        img.save(buf, format="JPEG", quality=quality)
        if buf.tell() <= max_bytes:
            break
        quality -= 10
    buf.seek(0)
    return buf.read()

# ── Función análisis con Claude ───────────────────────────────────────────────
def analizar_con_claude(fotos_list, ambiente, elemento, api_key):
    client = anthropic.Anthropic(api_key=api_key)

    content = []
    for i, f in enumerate(fotos_list):
        f.seek(0)
        raw = f.read()
        img_bytes = comprimir_imagen(raw)
        img_b64 = base64.standard_b64encode(img_bytes).decode("utf-8")
        tipo = "panorámica (contexto general del ambiente)" if i == 0 else f"detalle #{i} (acercamiento al defecto)"
        content.append({"type": "text", "text": f"Fotografía {i+1} — {tipo}:"})
        content.append({"type": "image", "source": {"type": "base64", "media_type": "image/jpeg", "data": img_b64}})

    content.append({"type": "text", "text": f"""
Eres un inspector técnico de inmuebles profesional en Perú. Analiza las fotografías anteriores.

Ambiente: {ambiente}
Elemento: {elemento if elemento else "identifícalo tú"}

Responde SOLO con JSON sin texto adicional ni backticks:
{{
  "elemento": "nombre exacto del elemento inspeccionado",
  "defecto": "descripción concisa del defecto en negritas",
  "accion": "acción de solución en negritas",
  "resultado": "resultado esperado sin negritas",
  "tolerancia": "Rechazado | Observado | Informativo | Validado | N/A"
}}

Criterios de tolerancia:
- Rechazado: daño grave, requiere sustitución completa
- Observado: defecto reparable, necesita corrección
- Informativo: dentro de tolerancias, se registra como respaldo
- Validado: cumple estándares, sin observaciones
- N/A: no aplicable

El campo "defecto" debe describir exactamente lo que se ve en las fotos.
El campo "accion" debe ser la solución técnica específica.
El campo "resultado" debe indicar el propósito de la corrección.
"""})

    response = client.messages.create(
        model="claude-opus-4-5",
        max_tokens=600,
        messages=[{"role": "user", "content": content}]
    )
    texto = response.content[0].text.strip().replace("```json","").replace("```","").strip()
    return json.loads(texto)

# ── Botón analizar ────────────────────────────────────────────────────────────
if st.button("🤖 Analizar fotos con IA", use_container_width=True):
    if not api_key:
        st.error("❌ Ingresa tu API Key en el panel izquierdo")
    elif not fotos:
        st.error("❌ Sube al menos una foto")
    else:
        with st.spinner("Claude está analizando las fotos..."):
            try:
                resultado = analizar_con_claude(fotos, ambiente_sel, elemento_input, api_key)
                st.session_state["analisis_temp"] = resultado
                st.session_state["fotos_temp"]    = fotos
                st.session_state["ambiente_temp"] = ambiente_sel
            except Exception as e:
                st.error(f"Error: {str(e)}")

# ── Edición y guardado ────────────────────────────────────────────────────────
if "analisis_temp" in st.session_state:
    st.markdown('<div class="titulo-morado">✏️ Revisa y edita la observación antes de guardar</div>',
                unsafe_allow_html=True)
    r = st.session_state["analisis_temp"]

    elemento_final  = st.text_input("Elemento inspeccionado", value=r.get("elemento",""), key="e_elem")
    
    st.markdown("**Observación** (formato: **Defecto** Requiere **acción** para resultado)")
    col_d, col_a, col_re = st.columns(3)
    with col_d:
        defecto_final  = st.text_input("**Defecto** (negrita)", value=r.get("defecto",""), key="e_def")
    with col_a:
        accion_final   = st.text_input("**Acción** (negrita)", value=r.get("accion",""), key="e_acc")
    with col_re:
        resultado_final = st.text_input("resultado esperado", value=r.get("resultado",""), key="e_res")

    tolerancia_final = st.selectbox("Tolerancia", TOLERANCIAS,
                                     index=TOLERANCIAS.index(r.get("tolerancia","Validado"))
                                     if r.get("tolerancia") in TOLERANCIAS else 3,
                                     key="e_tol")

    # Preview de observación
    st.markdown(f"**Vista previa:** **{defecto_final}** Requiere **{accion_final}** para {resultado_final}")

    if st.button("💾 Guardar observación", use_container_width=True):
        fotos_guardadas = []
        for f in st.session_state["fotos_temp"]:
            f.seek(0)
            fotos_guardadas.append({"nombre": f.name, "bytes": f.read()})

        st.session_state.observaciones.append({
            "n":          len(st.session_state.observaciones) + 1,
            "ambiente":   st.session_state["ambiente_temp"],
            "elemento":   elemento_final,
            "defecto":    defecto_final,
            "accion":     accion_final,
            "resultado":  resultado_final,
            "tolerancia": tolerancia_final,
            "fotos":      fotos_guardadas,
        })
        del st.session_state["analisis_temp"]
        del st.session_state["fotos_temp"]
        del st.session_state["ambiente_temp"]
        st.session_state.upload_key += 1
        st.success(f"✅ Observación #{len(st.session_state.observaciones)} guardada")
        st.rerun()

st.divider()

# ── OBSERVACIONES GUARDADAS ───────────────────────────────────────────────────
if st.session_state.observaciones:
    st.markdown('<div class="titulo-negro">📋 OBSERVACIONES REGISTRADAS</div>', unsafe_allow_html=True)

    # Contador por tolerancia
    resumen = {}
    for obs in st.session_state.observaciones:
        resumen[obs["tolerancia"]] = resumen.get(obs["tolerancia"], 0) + 1

    cols = st.columns(5)
    for i, tol in enumerate(TOLERANCIAS):
        badge = tol.replace("/","").replace(" ","")
        with cols[i]:
            st.markdown(f'<div class="badge-{badge}">{tol}: {resumen.get(tol,0)}</div>',
                        unsafe_allow_html=True)
    st.markdown("")

    # Por ambiente
    ambientes_usados = list(dict.fromkeys([o["ambiente"] for o in st.session_state.observaciones]))
    for amb in ambientes_usados:
        obs_amb = [o for o in st.session_state.observaciones if o["ambiente"] == amb]
        with st.expander(f"📐 {amb}  —  {len(obs_amb)} observación(es)", expanded=False):
            for obs in obs_amb:
                badge = obs["tolerancia"].replace("/","").replace(" ","")
                st.markdown(f'<div class="badge-{badge}">{obs["tolerancia"]}</div>', unsafe_allow_html=True)
                st.markdown(f"**#{obs['n']} — {obs['elemento']}**")
                st.markdown(f"**{obs['defecto']}** Requiere **{obs['accion']}** para {obs['resultado']}")
                # Miniaturas
                foto_cols = st.columns(min(len(obs["fotos"]), 5))
                for j, foto in enumerate(obs["fotos"]):
                    with foto_cols[j]:
                        st.image(foto["bytes"], width=100,
                                 caption="🌐" if j==0 else f"🔍{j}")
                st.markdown("---")

    if st.button("🗑️ Eliminar última observación"):
        st.session_state.observaciones.pop()
        st.rerun()

    st.divider()

    # ── EXPORTAR WORD ─────────────────────────────────────────────────────────
    st.markdown('<div class="titulo-negro">📄 EXPORTAR INFORME WORD</div>', unsafe_allow_html=True)

    COLORES_TOL = {
        "Rechazado":   RGBColor(0x1a,0x1a,0x1a),
        "Observado":   RGBColor(0xE0,0x7B,0x39),
        "Informativo": RGBColor(0x7B,0x5E,0xA7),
        "Validado":    RGBColor(0x9B,0x7F,0xBF),
        "N/A":         RGBColor(0x88,0x88,0x88),
    }
    HEX_TOL = {
        "Rechazado":"1a1a1a","Observado":"E07B39",
        "Informativo":"7B5EA7","Validado":"9B7FBF","N/A":"888888"
    }

    def set_shd(cell_or_para, hex_color):
        el = cell_or_para._tc if hasattr(cell_or_para,"_tc") else cell_or_para._p
        pr = el.get_or_add_tcPr() if hasattr(cell_or_para,"_tc") else el.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),  "clear")
        shd.set(qn("w:color"),"auto")
        shd.set(qn("w:fill"), hex_color)
        pr.append(shd)

    def titulo_fila(doc, texto, hex_bg, size=11):
        p = doc.add_paragraph()
        r = p.add_run(texto)
        r.bold = True; r.font.size = Pt(size)
        r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),hex_bg)
        p._p.get_or_add_pPr().append(shd)
        return p

    def generar_word():
        doc = Document()
        for sec in doc.sections:
            sec.top_margin = sec.bottom_margin = Cm(2)
            sec.left_margin = sec.right_margin = Cm(2.5)

        # Portada
        titulo_fila(doc, "INFORME DE INSPECCIÓN TÉCNICA DE INMUEBLE", "000000", 14)
        doc.add_paragraph("")

        # Datos del informe
        tabla_d = doc.add_table(rows=0, cols=2)
        tabla_d.style = "Table Grid"
        for lbl, val in [
            ("Propietario:", propietario), ("Teléfono:", telefono),
            ("Inmobiliaria:", inmobiliaria), ("Proyecto:", proyecto),
            ("Dirección:", direccion), ("N° Departamento:", nro_depto),
            ("Metraje:", f"{metraje} m²"), ("Fecha:", str(fecha_insp)),
            ("Inspector:", "HUALLPA BARZOLA, Alex — Ingeniero Civil"),
        ]:
            row = tabla_d.add_row()
            set_shd(row.cells[0], "E8E0FA")
            r = row.cells[0].paragraphs[0].add_run(lbl)
            r.bold = True; r.font.color.rgb = RGBColor(0x85,0x63,0xE1)
            row.cells[1].paragraphs[0].add_run(val or "")

        doc.add_paragraph("")

        # Leyenda
        titulo_fila(doc, "CATEGORÍAS DE TOLERANCIA", "000000", 10)
        tabla_l = doc.add_table(rows=0, cols=2)
        tabla_l.style = "Table Grid"
        for tol, desc in [
            ("Rechazado","Se requiere sustitución completa"),
            ("Observado","Defecto reparable, necesita corrección"),
            ("Informativo","Dentro de tolerancias, se registra como respaldo"),
            ("Validado","No presenta observaciones"),
            ("N/A","No aplicable al espacio"),
        ]:
            row = tabla_l.add_row()
            set_shd(row.cells[0], HEX_TOL[tol])
            r = row.cells[0].paragraphs[0].add_run(tol)
            r.bold = True; r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
            row.cells[1].paragraphs[0].add_run(desc)

        doc.add_paragraph("")

        # Observaciones por ambiente
        ambientes_doc = list(dict.fromkeys([o["ambiente"] for o in st.session_state.observaciones]))
        for amb in ambientes_doc:
            obs_amb = [o for o in st.session_state.observaciones if o["ambiente"] == amb]
            titulo_fila(doc, amb.upper(), "8563E1", 11)

            # Tabla: N° | Elemento | Imágenes | T | Observación
            tabla = doc.add_table(rows=1, cols=5)
            tabla.style = "Table Grid"

            # Anchos de columna (en DXA: 1440 = 1 pulgada)
            anchos = [400, 1400, 2800, 400, 3560]
            for i, celda in enumerate(tabla.rows[0].cells):
                celda._tc.get_or_add_tcPr()
                set_shd(celda, "8563E1")
                p = celda.paragraphs[0]
                r = p.add_run(["N°","Elemento Inspeccionado","Imagen de Referencia","T","Observación"][i])
                r.bold = True; r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
                r.font.size = Pt(9)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            for obs in obs_amb:
                row = tabla.add_row()

                # N°
                p0 = row.cells[0].paragraphs[0]
                p0.add_run(str(obs["n"])).bold = True
                p0.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Elemento
                row.cells[1].paragraphs[0].add_run(obs["elemento"]).font.size = Pt(9)

                # Imágenes — todas en la misma celda
                cell_img = row.cells[2]
                p_img = cell_img.paragraphs[0]
                for j, foto in enumerate(obs["fotos"]):
                    try:
                        img_stream = BytesIO(foto["bytes"])
                        run_img = p_img.add_run()
                        run_img.add_picture(img_stream, width=Inches(1.2))
                        # Pequeño label debajo
                        lbl_run = p_img.add_run(f"\n{'Panorámica' if j==0 else f'Detalle {j}'}   ")
                        lbl_run.font.size = Pt(7)
                        lbl_run.font.color.rgb = RGBColor(0x88,0x88,0x88)
                    except:
                        p_img.add_run(f"[foto {j+1}] ")

                # T (Tolerancia inicial)
                p_t = row.cells[3].paragraphs[0]
                r_t = p_t.add_run(obs["tolerancia"][0])  # primera letra
                r_t.bold = True
                r_t.font.color.rgb = COLORES_TOL.get(obs["tolerancia"], RGBColor(0,0,0))
                r_t.font.size = Pt(9)
                p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Observación con formato: **defecto** Requiere **accion** para resultado
                p_obs = row.cells[4].paragraphs[0]
                
                r1 = p_obs.add_run(obs["defecto"])
                r1.bold = True
                r1.font.size = Pt(9)
                
                r2 = p_obs.add_run(f". Requiere ")
                r2.bold = False
                r2.font.size = Pt(9)
                
                r3 = p_obs.add_run(obs["accion"])
                r3.bold = True
                r3.font.size = Pt(9)
                
                r4 = p_obs.add_run(f" para {obs['resultado']}")
                r4.bold = False
                r4.font.size = Pt(9)

            doc.add_paragraph("")

        # Resumen final
        titulo_fila(doc, "RESUMEN DE INSPECCIÓN", "000000", 11)
        tabla_res = doc.add_table(rows=0, cols=3)
        tabla_res.style = "Table Grid"
        row_h = tabla_res.add_row()
        for i, h in enumerate(["Tolerancia", "Cantidad", "% del total"]):
            set_shd(row_h.cells[i], "8563E1")
            r = row_h.cells[i].paragraphs[0].add_run(h)
            r.bold = True; r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

        total = len(st.session_state.observaciones)
        for tol in TOLERANCIAS:
            cant = sum(1 for o in st.session_state.observaciones if o["tolerancia"] == tol)
            if cant > 0:
                row = tabla_res.add_row()
                set_shd(row.cells[0], HEX_TOL[tol])
                r = row.cells[0].paragraphs[0].add_run(tol)
                r.bold = True; r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
                row.cells[1].paragraphs[0].add_run(str(cant))
                row.cells[2].paragraphs[0].add_run(f"{round(cant/total*100)}%")

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.getvalue()

    nombre_archivo = f"Informe_{nro_depto or 'inspeccion'}_{fecha_insp}.docx"
    st.download_button(
        "📄 Descargar Informe Word",
        data=generar_word(),
        file_name=nombre_archivo,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )

    if st.button("🔄 Nueva inspección (limpiar todo)", use_container_width=True):
        st.session_state.observaciones = []
        st.session_state.upload_key += 1
        st.rerun()

else:
    st.info("👆 Sube las fotos de una observación y presiona **Analizar con IA** para comenzar.")
